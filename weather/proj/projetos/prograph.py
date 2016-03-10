# -*- coding: utf-8 -*-
#!/usr/bin/env python

import pandas as pd
import django
from param.models import Unidade, Param, Limites
from projetos.models import Medicao, PtoMonit, Projeto, Campanha
from django.template import Context, Template
from json import dumps
import requests
import base64
import time
from PIL import Image
from StringIO import StringIO
from docx import Document
from docx.shared import Inches


def geraGraphico(_col_pontos, _col_campanha, _parametro,  _idleg, _idclasse):

    ID_LEGISLACAO = _idleg
    ID_CLASSE = _idclasse

    cores = [ '#73A374', '#b0ccb0', '#a8cdd7', '#c0beaf', '#cec597', '#e8b7b7', '#66A7B9', '#6a6755',
                '#B4E789', '#F4AABA', '#94AECD', '#C2C174', '#8B878F', '#C3A3D7',
                '#77B25A', '#E7ACBC', '#7E93A6', '#93B483','#BC8A97','#B6C6D6',
                '#D6C7A1','#9FA0A1','#4E763F','#495D71','#9F6270','#725E29','#5E5C57',
                '#B4E789', '#F4AABA', '#94AECD', '#C2C174', '#8B878F', '#C3A3D7',
                '#77B25A', '#E7ACBC', '#7E93A6', '#93B483','#BC8A97','#B6C6D6',
                '#D6C7A1','#9FA0A1','#4E763F','#495D71','#9F6270','#725E29','#5E5C57',
            ]

    template = """
{chart: {  width:800, plotBackgroundColor: '#E6E6E6'  },
title: { text: ' ' },
xAxis: [{ categories: {{categories|safe}} , crosshair: true, labels: {style: {fontWeight: 'bold'}}}],
yAxis: [{title:{ text: '{{ str_param }}',style: {fontWeight: 'bold'} }, labels: { style: {fontWeight: 'bold'} }}],
credits: { text: ' ', href: ' '},
plotOptions: { line: { lineWidth: 2, marker: { enabled: false}}},
series: [ {% for item in series %}
{ color:'{{item.color}}', borderColor:'#E6E6E6', data:{{item.data}}, name:'{{ item.name}}', type:'{{item.type}}'},
{% endfor %}]}
"""

    ptomonit = []
    campanha = []
    grade = {}


    """
    Seleciona no banco os registros a serem tratados
    """
    col0 = Medicao.\
          objects.\
          filter( Parametro_FK_id=_parametro).\
          values("Parametro_FK_id",
                 "Campanha_FK_id",
                 "PtoMonit_FK_id",
                 "vlrLbl",
                 "vlr").\
          order_by("Parametro_FK_id",
                   "Campanha_FK_id",
                   "PtoMonit_FK_id")

    col = []
    for item in col0:
        if item["PtoMonit_FK_id"] in _col_pontos and \
           item["Campanha_FK_id"] in _col_campanha:

            if '<' not in item['vlrLbl']:
                col.append(item)


    """
    Cria os pontos e as campanhas
    """
    for item in col:
        key = item["PtoMonit_FK_id"]
        if not key in ptomonit:
            ptomonit.append(key)

        key = item["Campanha_FK_id"]
        if not key in campanha:
            campanha.append(key)

    """
    Cria grade Campanhas X Pontos
    """
    for item in col:
        camp = item["Campanha_FK_id"]
        if camp in grade:
            linha = grade[camp]
        else:
            linha = [0 for itpto in ptomonit]

        pto = item["PtoMonit_FK_id"]
        index = ptomonit.index(pto)
        linha[index] = item["vlr"]

        grade[camp] = linha


    lim_min = {}
    lim_max = {}
    """ Determina Limites """
    col_limite  = Limites.objects.filter(parametro_FK_id = _parametro,
                                     legislacao_FK_id = ID_LEGISLACAO,
                                     classe_FK_id = ID_CLASSE)
    if col_limite:
        obj_limite = col_limite[0]
        if obj_limite.vlr_max != 0:
            linha = [ float(obj_limite.vlr_max) for it in ptomonit ]
            lim_max = { 'name':'VMP', 'type':'line', 'color':'red', 'data':linha }

        if obj_limite.vlr_min != 0:
            color = 'red' if lim_max  else 'blue'
            linha = [ float(obj_limite.vlr_min) for it in ptomonit ]
            lim_min = { 'name':'VMP', 'type':'line', 'color':color, 'data':linha }

    """
    Cria cria categorias
    """
    categories = []
    for item in ptomonit:
        ponto = PtoMonit.objects.get(pk=item).sigla
        categories.append('{0}'.format(ponto))



    """
    Cria  as series
    """
    series = []
    pos = 0
    for item in grade:
        obj_campanha = Campanha.objects.get(pk=item).nome
        dados = [ float(it) for it in grade[item] ]
        if sum(dados) != 0:
            registro = {'name': '{0}'.format(obj_campanha),
                        'type':'column',
                        'data':dados,
                        'color': cores[pos]}
            series.append(registro)
            pos += 1

    if len(series) == 0:
        return None

    """
    Adiciona limites se houver
    """
    if lim_max:
        series.append(lim_max)

    if lim_min:
        series.append(lim_min)


    """
    Texto contendo o parâmetro do gráfico
    """
    parametro = Param.objects.get(pk=_parametro)
    str_param = u'{0} ({1})'.format(parametro.nome , parametro.unidade_FK)

    """
    Renderiza o template
    """
    contexto = Context({    'categories' : categories,
                            'str_param': str_param,
                            'series':series})
    templ = Template(template)
    saida = templ.render(contexto)

    saida= '{"infile":"' + saida.encode('utf-8','ignore') + '","constr":"Chart"}'
    saida = saida.replace('\n','')

    return saida






def geraTabela(document, _col_pontos, _col_campanha, _idleg, _idclasse):

    query = Medicao.objects.values('Parametro_FK_id', 'PtoMonit_FK_id', 'Campanha_FK_id', 'id')

    col = []
    qs = []
    for item in query:
        if item["PtoMonit_FK_id"] in _col_pontos and \
           item["Campanha_FK_id"] in _col_campanha:

            reg = (item['Parametro_FK_id'], item['PtoMonit_FK_id'])
            if reg not in col:
                col.append(reg)
                qs.append(item)

    df = pd.DataFrame.from_records(qs)
    saida = pd.pivot_table(df, index=[ 'Parametro_FK_id'], columns=['PtoMonit_FK_id'], values=['id'], fill_value=0)

    db_param  =  list(Param.objects.all().values('id', 'nome', 'unidade_FK__sigla').order_by('nome'))
    db_pontos  = list(PtoMonit.objects.all().values('id', 'sigla').order_by('sigla'))

    colunas = [ item[1] for item in  list(saida.columns)]
    pontos = [ item for item in db_pontos if item['id'] in colunas ]
    param = [ item for item in db_param if item['id'] in list(saida.index)]

    linhas = []

    """Header"""
    linhaBranco = [ '' for item in colunas]

    for x in param:
        linhas.append(list(linhaBranco))

    col_datas = []
    pos_profund = []

    for key_param in saida.index:
        reg_param = [ item for item in param if item['id'] == key_param][0]
        pos_param = param.index(reg_param)

        for key_pto in saida.loc[key_param].index:
            reg_pto = [ item for item in pontos if item['id'] == key_pto[1] ][0]
            pos_pto = pontos.index(reg_pto)

            key_vlr = saida.loc[key_param].get(key_pto)
            if key_vlr == 0:
                vlr = 0
            else:
                if key_param == 840:
                    vlr = Medicao.objects.get(pk=key_vlr).data
                else:
                    vlr = Medicao.objects.get(pk=key_vlr).vlrLbl
            linhas[pos_param][pos_pto] = vlr


    pos = 0
    col_linhas = []
    saida = []
    for item in linhas:
        key = param[pos]['id']
        insere = True
        if param[pos]['id'] == 840:
            newlinha = [ '{0:02d}/{1:02d}/{2}'.format(it.day, it.month, it.year) for it in item]
            newlinha.insert(0, ('Data', '', '') )
            col_linhas.append(newlinha)

            newlinha = ['{0:02d}:{1:02d}'.format(it.hour, it.minute) for it in item]
            newlinha.insert(0, ('Hora', '', '') )
            col_linhas.append(newlinha)
            insere = False


        if param[pos]['id'] == 798:
            newlinha =list(item)
            newlinha.insert(0, ('Profundidade', 'm', '') )
            col_linhas.append(newlinha)
            insere = False

        if insere:
            reg = [ (it['nome'], it['unidade_FK__sigla'], '') for it in param if it['id']==key][0]
            item.insert(0, reg)
            saida.append(item)
        pos +=1

    for it in col_linhas:
       saida.insert(0,it)

    main_rel = [ item[1:] for item in saida ]
    left_rel = [ item[:1][0] for item in saida ]
    head_rel = [ item['sigla'] for item in pontos]

    step_col = 7
    start = 0
    total = len(main_rel[1])

    for start in range(0,total,step_col):
        final = min( start+step_col , total)
        table = document.add_table(rows=len(saida)+1, cols=(final-start)+3)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Ponto'
        hdr_cells[1].text = 'Unidade'
        hdr_cells[2].text = 'VMP'
        pos = 3
        for ptos in head_rel[start:final]:
            hdr_cells[pos].text = ptos
            pos += 1

        row = 0
        for item in main_rel:
           pos = 0
           row_cells = table.rows[row+1].cells
           for it in left_rel[row]:
               row_cells[pos].text = u'{0}'.format(it)
               pos += 1

           for coluna in item[start:final]:
               row_cells[pos].text = u'{0}'.format(coluna)
               pos += 1
           row += 1

        document.add_page_break()

































def processa(_col_pontos, _col_campanha, _col_parametro,  _idleg, _idclasse):

    document = Document()
    document.add_heading('Document Title', 0)

    URL_PHANTOM = 'http://10.3.0.29:3003'
    headers ={ 'Content-Type': 'application/json', '-X POST':'' }
    indice = 0

    indice += 1
    document.add_heading(u'Quadro 6.{0} - - Resultados do monitoramento da qualidade da água'.format(indice, ''), level=2)
    geraTabela(document, _col_pontos, _col_campanha, _idleg, _idclasse)

    for parametro  in _col_parametro:

        """ Não processa parametros Data e Profundidade"""
        if parametro in (798,840):
            continue

        saida = geraGraphico(_col_pontos, _col_campanha, parametro, _idleg, _idclasse)
        if not saida:
            continue
        r = requests.post(URL_PHANTOM, data=saida, headers=headers)
        png_recovered = base64.decodestring(r.content)
        path = "/tmp/{0}.png".format(int(time.time()*1000))
        f = open(path, "w")
        f.write(png_recovered)
        f.close()

        indice += 1
        obj_param = Param.objects.get(pk=parametro)
        nome = obj_param.nome
        texto = obj_param.texto
        document.add_heading(u'6.{0} - {1}'.format(indice, nome), level=1)

        indice += 1
        document.add_heading(u'Figura 6.{0} -Resultados obtidos para o parâmetro {1}'.format(indice, nome), level=2)
        document.add_paragraph(texto)
        document.add_picture(path, width=Inches(5.2))
        document.add_page_break()
    path = "/tmp/{0}.docx".format(int(time.time()*1000))
    document.save(path)
    return path


