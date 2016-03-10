import pandas as pd
from projetos.models import Medicao, PtoMonit
from param.models import Param
from PIL import Image
from StringIO import StringIO
from docx import Document
from docx.shared import Inches
from docx.shared import Pt

def run():
    qs1 = Medicao.objects.filter(Campanha_FK_id=15, PtoMonit_FK__parent_id=2235).values('Parametro_FK_id', 'PtoMonit_FK_id', 'id')

    col = []
    qs = []
    for item in qs1:
        """
        if item['PtoMonit_FK_id'] not in (2303,2287,2296):
            continue
        """
        reg = (item['Parametro_FK_id'], item['PtoMonit_FK_id'])
        if reg not in col:
            col.append(reg)
            qs.append(item)

    df = pd.DataFrame.from_records(qs)
    saida = pd.pivot_table(df, index=[ 'Parametro_FK_id'], columns=['PtoMonit_FK_id'], values=['id'], fill_value=0)

    db_param  =  list(Param.objects.all().values('id', 'nome', 'unidade_FK__sigla').order_by('nome'))
    db_pontos  = list(PtoMonit.objects.all().values('id', 'sigla').order_by('sigla'))

    colunas = [ item[1] for item in  list(saida.columns)]
    pontos = [ item for item in db_pontos if item['id'] in colunas ]
    param = [ item for item in db_param if item['id'] in list(saida.index)]

    linhas = []

    """Header"""
    linhaBranco = [ '' for item in colunas]

    for x in param:
        linhas.append(list(linhaBranco))

    col_datas = []
    pos_profund = []

    for key_param in saida.index:
        reg_param = [ item for item in param if item['id'] == key_param][0]
        pos_param = param.index(reg_param)

        for key_pto in saida.loc[key_param].index:
            reg_pto = [ item for item in pontos if item['id'] == key_pto[1] ][0]
            pos_pto = pontos.index(reg_pto)

            key_vlr = saida.loc[key_param].get(key_pto)
            if key_vlr == 0:
                vlr = 0
            else:
                if key_param == 840:
                    vlr = Medicao.objects.get(pk=key_vlr).data
                else:
                    vlr = Medicao.objects.get(pk=key_vlr).vlrLbl
            linhas[pos_param][pos_pto] = vlr


    pos = 0
    col_linhas = []
    saida = []
    for item in linhas:
        key = param[pos]['id']
        insere = True
        if param[pos]['id'] == 840:
            newlinha = [ '{0:02d}/{1:02d}/{2}'.format(it.day, it.month, it.year) for it in item]
            newlinha.insert(0, ('Data', '') )
            col_linhas.append(newlinha)

            newlinha = ['{0:02d}:{1:02d}'.format(it.hour, it.minute) for it in item]
            newlinha.insert(0, ('Hora', '') )
            col_linhas.append(newlinha)
            insere = False


        if param[pos]['id'] == 798:
            newlinha =list(item)
            newlinha.insert(0, ('Profundidade', 'm') )
            col_linhas.append(newlinha)
            insere = False

        if insere:
            reg = [ (it['nome'], it['unidade_FK__sigla']) for it in param if it['id']==key][0]
            item.insert(0, reg)
            saida.append(item)
        pos +=1

    for it in col_linhas:
       saida.insert(0,it)

    main_rel = [ item[1:] for item in saida ]
    left_rel = [ item[:1][0] for item in saida ]
    head_rel = [ item['sigla'] for item in pontos]

    document = Document()
    document.add_heading('Document Title', 0)

    step_cols = 7
    start = 0


    table = document.add_table(rows=len(saida)+2, cols=step_cols+2)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = ''
    hdr_cells[1].text = ''
    pos = 2
    for ptos in head_rel[start:step_cols]:
        hdr_cells[pos].text = ptos
        pos += 1

    row = 1
    for item in main_rel:
       pos = 0
       row_cells = table.rows[row].cells
       for it in left_rel[row-1]:
           row_cells[pos].text = u'{0}'.format(it)
           pos += 1

       for coluna in item[start:step_cols]:
           row_cells[pos].text = u'{0}'.format(coluna)
           pos += 1
       row += 1

    document.add_page_break()
    document.save('/home/wbeirigo/demo.docx')
