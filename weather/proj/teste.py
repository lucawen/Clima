# -*- coding: utf-8 -*-
#!/usr/bin/env python


from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import docx

document = Document('/home/wbeirigo/Clima/weather/proj/matriz.docx')

"""
style = document.styles['']
font = style.font
font.name = 'Arial'
font.size = Pt(9)


section = document.sections[0]
section.page_width  = Inches(8.267)
section.page_height  = Inches(11.692)
"""

for item in list(document.styles):
   it = u'{0}'.format(item)
   if 'docx.styles.style._CharacterStyle' in (it):
        print(item)
   else:
       print(it)



document.add_heading('first item in unordered list', level=1)
document.add_heading( 'first item in ordered list',level=2 )
document.add_heading( 'first item in ordered list',level=3 )

document.add_paragraph('yuyuy987y78y87y897y', style='Normal')

recordset = [
        {'qtd':1, 'id':1, 'desc':'uiuiuiu'},
        {'qtd':1, 'id':1, 'desc':'uiuiuiu'},
        {'qtd':1, 'id':1, 'desc':'uiuiuiu'},
        {'qtd':1, 'id':1, 'desc':'uiuiuiu'},
        {'qtd':1, 'id':1, 'desc':'uiuiuiu'},
        ]

table = document.add_table(rows=1, cols=3)
#table.style = 'TableGrid'
table.autofit = True
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'qtd'
hdr_cells[1].text = 'id'
hdr_cells[2].text = 'desc'
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item['qtd'])
    row_cells[1].text = str(item['id'])
    row_cells[2].text = item['desc']

document.add_page_break()

document.save('demo.docx')
